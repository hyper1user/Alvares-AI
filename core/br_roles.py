"""
Бізнес-логіка для генерації БР з розподілом по ролях.
"""
import os
import sys
from datetime import datetime, timedelta
from typing import List, Dict, Tuple, Optional

# Додаємо кореневу директорію проєкту в sys.path для імпортів
if getattr(sys, 'frozen', False):
    _PROJECT_ROOT = os.path.dirname(sys.executable)
else:
    _PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if _PROJECT_ROOT not in sys.path:
    sys.path.insert(0, _PROJECT_ROOT)

from br_updater import get_tabel_date, get_soldiers_from_tabel, _get_soldiers_from_tabel_detailed, pib_to_document_format, normalize_pib, get_soldiers_returning_from_rop
from br_calculator import get_br_number
from excel_processor import TabelReader
from data.database import (
    get_all_roles, get_role_composition, get_all_personnel,
    set_personnel_role, upsert_personnel_batch, get_connection
)

# Маппінг ролей -> плейсхолдерів у шаблоні Word
PLACEHOLDER_MAP = {
    "Заступник командира роти": "{{ROLE_ZKR}}",
    "Офіцер з МПЗ": "{{ROLE_PPP}}",
    "Старший технік роти": "{{ROLE_SENIOR_TECH}}",
    "Головний сержант роти": "{{ROLE_FIRST_SERGEANT}}",
    "Сержант із матеріального забезпечення": "{{ROLE_SUPPLY_SERGEANT}}",
    "Старший бойовий медик": "{{ROLE_MEDIC}}",
    "Група евакуації": "{{ROLE_EVAC_GROUP}}",
    "Водій групи евакуації": "{{ROLE_EVAC_DRIVER}}",
    "Екіпажі розрахунків БМП-1ЛБ": "{{ROLE_BMP_CREWS}}",
    "Командири штурмових взводів": "{{ROLE_VZVOD}}",
    "Водії роти": "{{ROLE_DRIVERS}}",
    "Чергові зв'язківці": "{{ROLE_SIGNAL}}",
    "Резервні групи": "{{ROLE_RESERVE}}",
    "Супровід FPV": "{{ROLE_FPV_ESCORT}}",
    "Водії логістики": "{{ROLE_LOGISTICS_DRIVERS}}",
}

# Дефолтні значення для ролей, якщо немає бійців з mark==100
ROLE_DEFAULTS = {
    "{{ROLE_SENIOR_TECH}}": "старший сержант ЮХТА Андрій Петрович",
    "{{ROLE_FIRST_SERGEANT}}": "молодший сержант ШТУНДЕР Антон Валерійович",
    "{{ROLE_SUPPLY_SERGEANT}}": "сержант КЛИМОВЕЦ Сергій Григорович",
    "{{ROLE_MEDIC}}": "молодший сержант БАРТ Володимир Григорович",
}

# Ролі, абзац яких видаляється якщо немає бійців
ROLE_REMOVE_IF_EMPTY = {"{{ROLE_PPP}}"}


def auto_assign_role(position: str) -> Optional[str]:
    """
    Визначає роль за ключовими словами у назві посади.
    Повертає назву ролі або None.
    """
    pos_lower = position.lower().strip()
    if not pos_lower:
        return None

    # Заступник командира роти
    if "заступник" in pos_lower and "командир" in pos_lower:
        return "Заступник командира роти"

    # Офіцер з МПЗ (морально-психологічне забезпечення)
    if "мпз" in pos_lower or "морально" in pos_lower and "псих" in pos_lower:
        return "Офіцер з МПЗ"

    # Медик
    if "медик" in pos_lower or "медичн" in pos_lower or "санітар" in pos_lower:
        return "Старший бойовий медик"

    # Командир взводу
    if "командир" in pos_lower and "взвод" in pos_lower:
        return "Командири штурмових взводів"

    has_evak = "евак" in pos_lower
    has_vodiy = "водій" in pos_lower or "водiй" in pos_lower  # обидва варіанти "і"

    # Пріоритетна логіка для евакуації/водіїв
    if has_evak and has_vodiy:
        return "Водій групи евакуації"
    if has_evak:
        return "Група евакуації"

    # Перевіряємо "головний сержант" ДО простого "водій"
    if "головний сержант" in pos_lower:
        return "Головний сержант роти"
    if "матеріаль" in pos_lower:
        return "Сержант із матеріального забезпечення"

    if has_vodiy:
        return "Водії роти"

    if "fpv" in pos_lower:
        return "Супровід FPV"
    if "логістик" in pos_lower:
        return "Водії логістики"

    keyword_map = [
        ("зв'яз", "Чергові зв'язківці"),
        ("зв'яз", "Чергові зв'язківці"),  # апостроф варіант
        ("технік", "Старший технік роти"),
        ("бмп", "Екіпажі розрахунків БМП-1ЛБ"),
    ]
    for keyword, role_name in keyword_map:
        if keyword in pos_lower:
            return role_name

    return None


def import_personnel_from_tabel(tabel_file: str, sheet_name: str) -> int:
    """
    Імпортує особовий склад з аркуша табеля до БД.
    Повертає кількість імпортованих записів.
    """
    reader = TabelReader(tabel_file)
    reader.load_workbook()
    soldiers = reader.read_month_data(sheet_name)

    records = []
    for s in soldiers:
        if s.pib and s.pib.strip():
            records.append((s.pib, s.rank, s.position))

    if records:
        return upsert_personnel_batch(records)
    return 0


def get_active_personnel_for_month(tabel_file: str, sheet_name: str) -> List[Dict]:
    """
    Повертає бійців з позначками "100" та/або "роп" за обраний місяць.
    Формат повернення збігається з get_all_personnel() (pib, rank, position, role_id, role_name).
    """
    reader = TabelReader(tabel_file)
    reader.load_workbook()
    soldiers = reader.read_month_data(sheet_name)

    active_pibs = set()
    for s in soldiers:
        if s.days_100_combined:
            active_pibs.add(s.pib.strip())

    all_personnel = get_all_personnel()
    return [p for p in all_personnel if p["pib"].strip() in active_pibs]


def auto_assign_all_roles() -> Dict[str, int]:
    """
    Автопризначає ролі для тих, хто ще не має призначеної ролі.
    Повертає {role_name: count}.
    """
    personnel = get_all_personnel()
    roles = get_all_roles()
    role_name_to_id = {name: rid for rid, name in roles}

    stats = {}
    for person in personnel:
        if person["role_id"] is not None:
            continue  # вже має роль
        role_name = auto_assign_role(person["position"])
        if role_name and role_name in role_name_to_id:
            set_personnel_role(person["pib"], role_name_to_id[role_name])
            stats[role_name] = stats.get(role_name, 0) + 1

    return stats


def get_soldiers_100_for_br_date(tabel_file: str, br_date: datetime) -> List[Tuple[str, str]]:
    """
    Для дати БР обчислює дату табеля (BR+1), читає табель,
    повертає бійців з mark==100 АБО mark=="роп" як [(pib, rank)].
    Також додає тих, у кого роп щойно закінчився (вчора="роп", сьогодні=порожньо).
    """
    tabel_date = get_tabel_date(br_date)
    soldiers_100, soldiers_rop, _ = _get_soldiers_from_tabel_detailed(tabel_file, tabel_date)

    result = soldiers_100 + soldiers_rop

    # Бійці, у яких роп закінчився і клітинка сьогодні порожня
    returning = get_soldiers_returning_from_rop(tabel_file, br_date)
    if returning:
        existing = {normalize_pib(p) for p, _ in result}
        result += [(p, r) for p, r in returning if normalize_pib(p) not in existing]

    return result


def build_composition_for_date(
    tabel_file: str, br_date: datetime
) -> Dict[str, List[Dict]]:
    """
    Формує склад БР на дату: для кожної ролі — список людей з mark==100 або mark=="роп",
    а також тих, у кого роп щойно закінчився (повернулись з позиції).
    Ті, хто без ролі → "Резервні групи".
    """
    soldiers_100 = get_soldiers_100_for_br_date(tabel_file, br_date)

    # Нормалізовані ПІБ для порівняння
    pibs_100_map = {}  # normalized_pib -> (original_pib, rank)
    for pib, rank in soldiers_100:
        norm = normalize_pib(pib)
        pibs_100_map[norm] = (pib, rank)

    role_composition = get_role_composition()
    roles = get_all_roles()

    result = {}
    for _, role_name in roles:
        result[role_name] = []

    assigned_pibs = set()

    for role_name, members in role_composition.items():
        for member in members:
            norm = normalize_pib(member["pib"])
            if norm in pibs_100_map:
                orig_pib, orig_rank = pibs_100_map[norm]
                result[role_name].append({
                    "pib": orig_pib,
                    "rank": orig_rank or member["rank"],
                    "position": member["position"]
                })
                assigned_pibs.add(norm)

    # Бійці з 100 без ролі — НЕ включаються до БР

    return result


def _remove_paragraph(paragraph):
    """Видаляє параграф з документа."""
    p_element = paragraph._element
    parent = p_element.getparent()
    if parent is not None:
        parent.remove(p_element)


def _replace_in_paragraph(paragraph, key: str, value: str, size_pt: int = 10):
    """Замінює плейсхолдер у параграфі, підтримує багаторядкові значення."""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    import copy

    full_text = paragraph.text
    if key not in full_text:
        return

    # Примусово Times New Roman з вказаним розміром
    font_name = "Times New Roman"
    font_size = Pt(size_pt)
    font_bold = None
    if paragraph.runs:
        font_bold = paragraph.runs[0].font.bold

    new_text = full_text.replace(key, value)

    lines = new_text.split("\n")

    if len(lines) <= 1:
        # Однорядковий текст — просто замінюємо runs
        p_element = paragraph._element
        for r in list(p_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')):
            p_element.remove(r)
        run = paragraph.add_run(new_text)
        run.font.name = font_name
        run.font.size = font_size
        if font_bold is not None:
            run.font.bold = font_bold
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = run._element.makeelement(qn('w:rFonts'), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        return

    # Багаторядковий текст — створюємо окремі абзаци (звичайний Enter)
    p_element = paragraph._element
    parent = p_element.getparent()
    p_index = list(parent).index(p_element)

    # Зберігаємо pPr (властивості абзацу) для копіювання
    pPr = p_element.find(qn('w:pPr'))

    def _make_paragraph_with_text(text):
        """Створює новий абзац з текстом та форматуванням."""
        new_p = p_element.makeelement(qn('w:p'), {})
        if pPr is not None:
            new_p.append(copy.deepcopy(pPr))
        r = new_p.makeelement(qn('w:r'), {})
        rPr_new = r.makeelement(qn('w:rPr'), {})
        rFonts_new = r.makeelement(qn('w:rFonts'), {
            qn('w:ascii'): font_name, qn('w:hAnsi'): font_name,
            qn('w:cs'): font_name, qn('w:eastAsia'): font_name,
        })
        rPr_new.append(rFonts_new)
        half_pts = str(size_pt * 2)
        sz = r.makeelement(qn('w:sz'), {qn('w:val'): half_pts})
        rPr_new.append(sz)
        szCs = r.makeelement(qn('w:szCs'), {qn('w:val'): half_pts})
        rPr_new.append(szCs)
        if font_bold:
            b = r.makeelement(qn('w:b'), {})
            rPr_new.append(b)
        r.append(rPr_new)
        t = r.makeelement(qn('w:t'), {})
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        new_p.append(r)
        return new_p

    # Перший рядок залишається в оригінальному абзаці
    for r in list(p_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')):
        p_element.remove(r)
    run = paragraph.add_run(lines[0])
    run.font.name = font_name
    run.font.size = font_size
    if font_bold is not None:
        run.font.bold = font_bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)

    # Решта рядків — нові абзаци після оригінального
    for j, line in enumerate(lines[1:], start=1):
        new_p = _make_paragraph_with_text(line)
        parent.insert(p_index + j, new_p)


def _make_ack_run(p_elem, text, font_name='Times New Roman', size_half_pt='24'):
    """Створює run з текстом та шрифтом Times New Roman 12."""
    from docx.oxml.ns import qn
    r = p_elem.makeelement(qn('w:r'), {})
    rPr = r.makeelement(qn('w:rPr'), {})
    rFonts = r.makeelement(qn('w:rFonts'), {
        qn('w:ascii'): font_name, qn('w:hAnsi'): font_name,
        qn('w:cs'): font_name, qn('w:eastAsia'): font_name,
    })
    rPr.append(rFonts)
    sz = r.makeelement(qn('w:sz'), {qn('w:val'): size_half_pt})
    rPr.append(sz)
    szCs = r.makeelement(qn('w:szCs'), {qn('w:val'): size_half_pt})
    rPr.append(szCs)
    r.append(rPr)
    t = r.makeelement(qn('w:t'), {})
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    return r


def _make_tab_run(p_elem):
    """Створює run з символом табуляції."""
    from docx.oxml.ns import qn
    r = p_elem.makeelement(qn('w:r'), {})
    tab = r.makeelement(qn('w:tab'), {})
    r.append(tab)
    return r


def _insert_ack_list(paragraph, members: list):
    """
    Замінює параграф з {{ACK_LIST}} на набір окремих параграфів:
    звання (ліворуч) | ____ підпис (центр) | Ім'я ПРІЗВИЩЕ (праворуч)
    Використовує tab stops для вирівнювання.
    """
    from docx.oxml.ns import qn
    from br_updater import pib_to_table_format

    parent = paragraph._element.getparent()
    ref = paragraph._element

    # Tab stop positions (in twips: 1440 twips = 1 inch)
    # center tab ~4500 twips (≈7.9cm), right tab ~9600 twips (≈16.9cm)
    center_pos = "4500"
    right_pos = "9600"

    for m in members:
        new_p = paragraph._element.makeelement(qn('w:p'), {})

        # Paragraph properties: tab stops + left alignment
        pPr = new_p.makeelement(qn('w:pPr'), {})
        tabs = new_p.makeelement(qn('w:tabs'), {})
        tab_center = new_p.makeelement(qn('w:tab'), {
            qn('w:val'): 'center', qn('w:pos'): center_pos, qn('w:leader'): 'none'
        })
        tab_right = new_p.makeelement(qn('w:tab'), {
            qn('w:val'): 'right', qn('w:pos'): right_pos, qn('w:leader'): 'none'
        })
        tabs.append(tab_center)
        tabs.append(tab_right)
        pPr.append(tabs)
        jc = new_p.makeelement(qn('w:jc'), {qn('w:val'): 'left'})
        pPr.append(jc)
        # Міжрядковий інтервал — одинарний
        spacing = new_p.makeelement(qn('w:spacing'), {
            qn('w:after'): '0', qn('w:line'): '240', qn('w:lineRule'): 'auto'
        })
        pPr.append(spacing)
        new_p.append(pPr)

        # 1) Звання (ліворуч)
        new_p.append(_make_ack_run(new_p, m['rank']))
        # Tab до центру
        new_p.append(_make_tab_run(new_p))
        # 2) Лінія для підпису (по центру)
        new_p.append(_make_ack_run(new_p, '____________________'))
        # Tab до правого краю
        new_p.append(_make_tab_run(new_p))
        # 3) Ім'я ПРІЗВИЩЕ (праворуч)
        new_p.append(_make_ack_run(new_p, pib_to_table_format(m['pib'])))

        parent.insert(list(parent).index(ref) + 1, new_p)
        ref = new_p

    # Видаляємо оригінальний параграф з {{ACK_LIST}}
    parent.remove(paragraph._element)


def get_br_from_4shb(br_4shb_file: str, tabel_date: datetime) -> Tuple[str, str]:
    """
    Знаходить номер та дату БР з файлу BR_4ShB.xlsx для дати табеля.
    Логіка "на завтра": шукаємо запис де дата = tabel_date.

    Args:
        br_4shb_file: Шлях до BR_4ShB.xlsx
        tabel_date: Дата табеля (BR+1)

    Returns:
        (номер_бр, дата_бр_форматована) або ("—", "—") якщо не знайдено
    """
    import openpyxl

    if not os.path.exists(br_4shb_file):
        return "—", "—"

    wb = openpyxl.load_workbook(br_4shb_file, data_only=True)
    ws = wb[wb.sheetnames[0]]

    target_date = tabel_date.date() if hasattr(tabel_date, 'date') else tabel_date

    # Шукаємо останній запис з потрібною датою (якщо декілька на одну дату)
    found_id = None
    found_date = None

    for row in range(2, ws.max_row + 1):
        cell_id = ws.cell(row, 1).value
        cell_date = ws.cell(row, 2).value

        if not cell_id or not cell_date:
            continue

        if hasattr(cell_date, 'date'):
            row_date = cell_date.date()
        elif isinstance(cell_date, str):
            from datetime import datetime as dt
            row_date = dt.strptime(cell_date[:10], "%Y-%m-%d").date()
        else:
            continue

        if row_date == target_date:
            found_id = str(cell_id)
            found_date = row_date

    if found_id and found_date:
        date_str = found_date.strftime("%d.%m.%Y")
        return found_id, date_str

    return "—", "—"


def _process_if_rop_block(doc, has_rop: bool):
    """
    Обробляє IF-блок {{IF_ROP}}...{{/IF_ROP}} у документі.
    Якщо has_rop=True — видаляє тільки маркери, залишає вміст.
    Якщо has_rop=False — видаляє весь блок (всі параграфи між маркерами включно).
    """
    paragraphs = doc.paragraphs
    start_idx = None
    end_idx = None

    for i, p in enumerate(paragraphs):
        if "{{IF_ROP}}" in p.text:
            start_idx = i
        if "{{/IF_ROP}}" in p.text:
            end_idx = i
            break

    if start_idx is None or end_idx is None:
        return

    if has_rop:
        # Видаляємо тільки маркери з тексту параграфів
        start_p = paragraphs[start_idx]
        end_p = paragraphs[end_idx]
        if start_p.text.strip() == "{{IF_ROP}}":
            _remove_paragraph(start_p)
        else:
            _replace_in_paragraph(start_p, "{{IF_ROP}}", "", size_pt=10)
        # Після видалення start_p індекси зсунулися, шукаємо end заново
        for p in doc.paragraphs:
            if "{{/IF_ROP}}" in p.text:
                if p.text.strip() == "{{/IF_ROP}}":
                    _remove_paragraph(p)
                else:
                    _replace_in_paragraph(p, "{{/IF_ROP}}", "", size_pt=10)
                break
    else:
        # Видаляємо весь блок: від start_idx до end_idx включно
        to_remove = list(paragraphs[start_idx:end_idx + 1])
        for p in to_remove:
            _remove_paragraph(p)


def generate_br_word(
    br_date: datetime,
    composition: Dict[str, List[Dict]],
    template_path: str,
    output_dir: str = "output",
    br_4shb_file: str = None,
    tabel_file: str = None,
    rop_txt_path: str = None,
    dodatky_path: str = None
) -> str:
    """
    Генерує Word-документ БР з шаблону, замінюючи плейсхолдери.
    Повертає шлях до створеного файлу.
    """
    from docx import Document

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Шаблон не знайдено: {template_path}")

    doc = Document(template_path)

    tabel_date = get_tabel_date(br_date)
    br_number = get_br_number(tabel_date)
    date_str = br_date.strftime("%d.%m.%Y")

    # Номер БР з BR_4ShB.xlsx (якщо файл вказано)
    # Шукаємо по даті БР (br_date), а не по даті табеля (tabel_date)
    if br_4shb_file:
        br_4shb_num, _ = get_br_from_4shb(br_4shb_file, br_date)
    else:
        br_4shb_num = br_number

    # Дата виконання = наступний день після дати БР (тобто tabel_date)
    execution_date_str = tabel_date.strftime("%d.%m.%Y")

    # Номер БР для шапки = порядковий номер дня року (від дати табеля)
    day_of_year = tabel_date.timetuple().tm_yday

    # Формуємо словник замін
    # {{бр}} = номер з BR_4ShB, {{дата_бр}} = дата самого БР
    # <<№*>> = порядковий номер дня року
    replacements = {
        "{{бр}}": br_4shb_num,
        "{{дата_бр}}": date_str,
        "<<Дата_виконання>>": execution_date_str,
        "<<№*>>": f"№{day_of_year}",
        "<<від 01.01.2026 р.>>": f"від {date_str} р.",
    }

    # Збираємо ПІБ бійців з РОП (перший день + продовження) для виключення з ролей
    rop_pibs = set()
    first_rop_entries = []
    continuing_rop = []
    if tabel_file:
        from br_updater import get_first_rop_entries, get_continuing_rop_entries
        first_rop_entries = get_first_rop_entries(tabel_file, br_date)
        continuing_rop = get_continuing_rop_entries(tabel_file, br_date)
        for pib, _, _ in first_rop_entries:
            rop_pibs.add(normalize_pib(pib))
        for pib, _, _ in continuing_rop:
            rop_pibs.add(normalize_pib(pib))

    for role_name, placeholder in PLACEHOLDER_MAP.items():
        members = composition.get(role_name, [])
        # Виключаємо бійців, які є в РОП — вони вказуються тільки в {{ROP_FIRST}} або {{ROP}}
        members = [m for m in members if normalize_pib(m["pib"]) not in rop_pibs]
        if members:
            parts = [pib_to_document_format(m["pib"], m["rank"]) for m in members]
            replacements[placeholder] = ", ".join(parts)
        else:
            if placeholder in ROLE_DEFAULTS:
                replacements[placeholder] = ROLE_DEFAULTS[placeholder]
            elif placeholder in ROLE_REMOVE_IF_EMPTY:
                replacements[placeholder] = None  # маркер для видалення абзацу
            else:
                replacements[placeholder] = "—"

    # ROP — бійці з продовженням "роп" (2-й+ день), перелік через кому
    # Завдання {{ROP1}}-{{ROP4}} з ROP.txt
    rop_placeholders = {}
    if continuing_rop:
        rop_list = ", ".join(
            pib_to_document_format(pib, rank)
            for pib, rank, pos in continuing_rop
        )
        replacements["{{ROP}}"] = rop_list
    else:
        replacements["{{ROP}}"] = None

    # Завдання з ROP.txt: {{ROP1}}-{{ROP4}}
    if rop_txt_path and os.path.exists(rop_txt_path):
        import re
        with open(rop_txt_path, "r", encoding="utf-8") as f:
            rop_content = f.read()
        for m in re.finditer(r"\{\{ROP(\d+)\}\}\s*(.+)", rop_content):
            marker = f"{{{{ROP{m.group(1)}}}}}"
            task_text = m.group(2).strip()
            rop_placeholders[marker] = task_text
    # Додаємо ROP-завдання до replacements (або None для видалення)
    for i in range(1, 5):
        marker = f"{{{{ROP{i}}}}}"
        replacements[marker] = rop_placeholders.get(marker, None)

    # Маркери з Dodatky.md: {{населений_пункт}}, {{КСП_РОТИ}}
    if dodatky_path:
        from core.dodatky_parser import get_location_for_date
        location = get_location_for_date(dodatky_path, br_date)
        replacements["{{населений_пункт}}"] = location["населений_пункт"]
        replacements["{{КСП_РОТИ}}"] = location["КСП_РОТИ"]
    else:
        replacements["{{населений_пункт}}"] = "—"
        replacements["{{КСП_РОТИ}}"] = "—"

    # {{ROP_FIRST}} — бійці з першим днем "роп"
    if first_rop_entries:
        rop_first_lines = ";\n".join(
            f"{pib_to_document_format(pib, rank)}, {pos}"
            for pib, rank, pos in first_rop_entries
        ) + ";"
        replacements["{{ROP_FIRST}}"] = rop_first_lines
    else:
        replacements["{{ROP_FIRST}}"] = ""

    # ACK_LIST — аркуш доведення: окремі параграфи для кожної людини
    from br_updater import pib_to_table_format
    all_members = []
    for members in composition.values():
        all_members.extend(members)
    # Додаємо бійців з ROP_FIRST до аркуша доведення
    for pib, rank, _ in first_rop_entries:
        all_members.append({"pib": pib, "rank": rank})
    ack_members = all_members

    # Для звичайних плейсхолдерів ставимо заглушку (буде замінено нижче)
    replacements["{{ACK_LIST}}"] = "—" if not ack_members else ""

    # Обробка IF-блоку {{IF_ROP}}...{{/IF_ROP}}
    _process_if_rop_block(doc, bool(first_rop_entries))

    # Замінюємо у параграфах (шрифт 10pt)
    paragraphs_to_remove = []
    for paragraph in doc.paragraphs:
        # Спеціальна обробка ACK_LIST — кожна людина як окремий параграф
        if "{{ACK_LIST}}" in paragraph.text and ack_members:
            _insert_ack_list(paragraph, ack_members)
            continue
        for key, value in replacements.items():
            if key in paragraph.text:
                if value is None:
                    paragraphs_to_remove.append(paragraph)
                    break
                _replace_in_paragraph(paragraph, key, value, size_pt=10)
    for p in paragraphs_to_remove:
        _remove_paragraph(p)

    # Замінюємо в таблицях (шрифт 10pt)
    paragraphs_to_remove = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            if value is None:
                                paragraphs_to_remove.append(paragraph)
                                break
                            _replace_in_paragraph(paragraph, key, value, size_pt=10)
    for p in paragraphs_to_remove:
        _remove_paragraph(p)

    os.makedirs(output_dir, exist_ok=True)
    output_file = os.path.join(output_dir, f"БР_ком_12шр_№{day_of_year}_від_{date_str.replace('.', '_')}.docx")
    doc.save(output_file)
    return output_file




