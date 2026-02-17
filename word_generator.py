"""
Модуль генерації Word-рапортів на підтвердження
"""
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import List
from excel_processor import SoldierData

class WordReportGenerator:
    """Клас для генерації Word-рапортів"""
    
    def __init__(self, template_file: str = None):
        """
        Args:
            template_file: Шлях до шаблону Word (якщо є)
        """
        self.template_file = template_file
    
    def create_confirmation_report(self, soldiers: List[SoldierData], month_name: str, 
                                 category: str, output_file: str) -> str:
        """
        Створює рапорт на підтвердження
        
        Args:
            soldiers: Список бійців
            month_name: Назва місяця (наприклад, "травень 2025")
            category: "100" або "30"
            output_file: Шлях до файлу для збереження
            
        Returns:
            str: Шлях до створеного файлу
        """
        
        # Створюємо новий документ або використовуємо шаблон
        if self.template_file and self.template_file.endswith('.docx'):
            doc = Document(self.template_file)
        else:
            doc = Document()
        
        # Заголовок документа
        self._add_header(doc, month_name, category)
        
        # Таблиця з даними
        self._add_soldiers_table(doc, soldiers, category)
        
        # Підпис
        self._add_signature(doc)
        
        # Зберігаємо документ
        doc.save(output_file)
        print(f"Створено рапорт: {output_file}")
        return output_file
    
    def _add_header(self, doc: Document, month_name: str, category: str):
        """Додає заголовок документа"""
        
        # Основний заголовок
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("ВІДОМІСТЬ")
        title_run.font.size = Inches(0.2)
        title_run.bold = True
        
        # Підзаголовок
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run("про участь військовослужбовців 12ШР 4 ШБ у бойових діях")
        
        # Назва підрозділу та період
        period_text = f"за {month_name} місяць"
        period = doc.add_paragraph()
        period.alignment = WD_ALIGN_PARAGRAPH.CENTER
        period.add_run(period_text)
        
        # Додаємо порожній рядок
        doc.add_paragraph()
        
        # Пояснення категорії
        explanation = doc.add_paragraph()
        if category == "100":
            explanation.add_run("1. Військовослужбовці, які безпосередньо брали участь у бойових діях")
        else:
            explanation.add_run("2. Військовослужбовці (забезпечуючі) військові частини в районі проведення бойових дій (забезпечувальні) в районі проведення бойових дій")
        
        doc.add_paragraph()
    
    def _add_soldiers_table(self, doc: Document, soldiers: List[SoldierData], category: str):
        """Додає таблицю з даними бійців"""
        
        # Створюємо таблицю
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        
        # Заголовки таблиці
        headers = [
            "№ п/п",
            "Військове звання",
            "Прізвище ім'я по батькові",
            "Період участі",
            "Підстава, № розпорядження, дата",
            "Категорія нарахувань",
            "Примітка"
        ]
        
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            # Робимо заголовки жирними
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Додаємо дані бійців
        for i, soldier in enumerate(soldiers, 1):
            row = table.add_row()
            
            # № п/п
            row.cells[0].text = str(i)
            
            # Звання
            row.cells[1].text = soldier.rank
            
            # ПІБ
            row.cells[2].text = soldier.pib
            
            # Період
            if category == "100" and soldier.days_100:
                period = self._get_period_string(soldier.days_100)
            elif category == "30" and soldier.days_30:
                period = self._get_period_string(soldier.days_30)
            else:
                period = ""
            row.cells[3].text = period
            
            # Підстава (номери БР/БН)
            if category == "100":
                br_list = soldier.get_br_list_100()
            else:
                br_list = soldier.get_br_list_30()
            row.cells[4].text = br_list
            
            # Сума (заглушка)
            if category == "100":
                row.cells[5].text = "100 000"
            else:
                row.cells[5].text = "30 000"
            
            # Примітка
            if soldier.has_no_payment_note():
                row.cells[6].text = "не виплачувати"
            else:
                row.cells[6].text = ""
    
    def _get_period_string(self, dates: List[datetime]) -> str:
        """Повертає рядок періоду"""
        if not dates:
            return ""
        
        dates.sort()
        start_date = dates[0]
        end_date = dates[-1]
        
        if start_date == end_date:
            return start_date.strftime("%d.%m.%Y")
        else:
            return f"з {start_date.strftime('%d.%m.%Y')} по {end_date.strftime('%d.%m.%Y')}"
    
    def _add_signature(self, doc: Document):
        """Додає підпис"""
        doc.add_paragraph()
        
        signature = doc.add_paragraph()
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        signature.add_run("Командир 12 штурмової роти 4 штурмового батальйону")
        
        # Порожній рядок для підпису
        doc.add_paragraph()
        
        signature_name = doc.add_paragraph()
        signature_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        signature_name.add_run("капітан _________________ Євген КРАСНИЙ")
        
        doc.add_paragraph()
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.add_run(f"«___» ___________ 2025 р.")

def create_sample_word_report():
    """Створює зразок Word-рапорту для тестування"""
    
    from excel_processor import TabelReader
    
    # Читаємо дані
    reader = TabelReader("Табель_Багатомісячний.xlsx")
    soldiers = reader.read_month_data("Травень_2025")
    
    # Фільтруємо бійців
    soldiers_100 = reader.get_soldiers_by_category(soldiers, "100", include_no_payment=True)
    soldiers_30 = reader.get_soldiers_by_category(soldiers, "30", include_no_payment=True)
    
    # Генеруємо рапорти
    generator = WordReportGenerator()
    
    # Рапорт на підтвердження 100к
    if soldiers_100:
        generator.create_confirmation_report(
            soldiers_100, 
            "травень 2025", 
            "100", 
            "Підтвердження_100к_травень_2025.docx"
        )
    
    # Рапорт на підтвердження 30к
    if soldiers_30:
        generator.create_confirmation_report(
            soldiers_30, 
            "травень 2025", 
            "30", 
            "Підтвердження_30к_травень_2025.docx"
        )

if __name__ == "__main__":
    create_sample_word_report()

